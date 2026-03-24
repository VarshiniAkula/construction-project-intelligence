"""Master document ingestion pipeline task."""
import os
import io
import csv
import json
import uuid
import logging
import httpx
import fitz  # PyMuPDF

from worker.celery_app import celery_app
from worker.shared.db import get_session
from worker.shared.minio_client import download_bytes, upload_bytes
from worker.shared.qdrant_client import ensure_collection, upsert_points

from qdrant_client import models as qdrant_models

# Import DB models from the API app
import sys
sys.path.insert(0, os.environ.get("API_APP_PATH", "/app/api_app"))

logger = logging.getLogger(__name__)


def _update_status(session, doc_id: str, status: str, error: str | None = None):
    """Update document status in database."""
    session.execute(
        session.get_bind().execute(
            f"UPDATE documents SET status = :status, processing_error = :error WHERE id = :id",
            {"status": status, "error": error, "id": doc_id},
        )
    )
    session.commit()


@celery_app.task(name="worker.tasks.ingest.ingest_document", bind=True, max_retries=2)
def ingest_document(self, document_id: str):
    """Full ingestion pipeline for a document."""
    session = get_session()

    try:
        # Load document record
        from sqlalchemy import text
        result = session.execute(
            text("SELECT * FROM documents WHERE id = :id"),
            {"id": document_id},
        )
        doc_row = result.mappings().first()
        if not doc_row:
            logger.error(f"Document {document_id} not found")
            return

        project_id = doc_row["project_id"]
        storage_key = doc_row["storage_key"]
        file_type = doc_row["file_type"]
        doc_type = doc_row["doc_type"]
        visibility_scope = doc_row["visibility_scope"]
        trade_scope = doc_row["trade_scope"] or ""

        logger.info(f"Starting ingestion for document {document_id} ({file_type})")

        # Update status
        session.execute(
            text("UPDATE documents SET status = 'rendering_pages' WHERE id = :id"),
            {"id": document_id},
        )
        session.commit()

        # Step 1: Download file from MinIO
        file_data = download_bytes(storage_key)

        if file_type == "pdf":
            pages_data = _process_pdf(session, document_id, project_id, file_data)
        elif file_type in ("png", "jpg", "jpeg"):
            pages_data = _process_image(session, document_id, project_id, file_data, file_type)
        elif file_type == "docx":
            pages_data = _process_docx(session, document_id, project_id, file_data)
        elif file_type == "xlsx":
            pages_data = _process_xlsx(session, document_id, project_id, file_data)
        elif file_type == "csv":
            pages_data = _process_csv(session, document_id, project_id, file_data)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Update page count
        session.execute(
            text("UPDATE documents SET page_count = :count, status = 'chunking' WHERE id = :id"),
            {"count": len(pages_data), "id": document_id},
        )
        session.commit()

        # Step 4: Chunk the content
        chunks = _chunk_pages(pages_data, doc_type)
        logger.info(f"Created {len(chunks)} chunks for document {document_id}")

        # Step 5: Generate embeddings and upsert to Qdrant
        session.execute(
            text("UPDATE documents SET status = 'embedding' WHERE id = :id"),
            {"id": document_id},
        )
        session.commit()

        if chunks:
            _embed_and_store(
                session, document_id, project_id, chunks,
                visibility_scope, trade_scope, doc_type,
            )

        # Mark as ready
        session.execute(
            text("UPDATE documents SET status = 'ready' WHERE id = :id"),
            {"id": document_id},
        )
        session.commit()
        logger.info(f"Document {document_id} ingestion complete")

    except Exception as e:
        logger.error(f"Ingestion failed for {document_id}: {e}")
        try:
            session.execute(
                text("UPDATE documents SET status = 'error', processing_error = :err WHERE id = :id"),
                {"err": str(e)[:1000], "id": document_id},
            )
            session.commit()
        except Exception:
            session.rollback()
        raise
    finally:
        session.close()


def _process_pdf(session, document_id: str, project_id: str, file_data: bytes) -> list[dict]:
    """Render PDF pages, extract text, create page records."""
    from sqlalchemy import text

    doc = fitz.open(stream=file_data, filetype="pdf")
    pages_data = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)

        # Extract text
        raw_text = page.get_text("text")

        # Render page image
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")

        # Upload page image to MinIO
        image_key = f"projects/{project_id}/documents/{document_id}/pages/{page_num + 1}.png"
        upload_bytes(image_key, img_bytes, content_type="image/png")

        # VLM analysis (optional, if endpoint available)
        vlm_result = _try_vlm_analysis(img_bytes, raw_text)

        page_summary = vlm_result.get("summary", raw_text[:500] if raw_text else "")
        cleaned_text = raw_text.strip()
        if vlm_result.get("summary") and not cleaned_text:
            cleaned_text = vlm_result["summary"]

        # Create page record
        page_id = str(uuid.uuid4())
        session.execute(
            text("""INSERT INTO document_pages
                (id, document_id, page_number, raw_text, cleaned_text, page_summary,
                 image_storage_key, extracted_json, created_at, updated_at)
                VALUES (:id, :doc_id, :pn, :raw, :cleaned, :summary, :img_key, :json,
                        NOW(), NOW())"""),
            {
                "id": page_id, "doc_id": document_id, "pn": page_num + 1,
                "raw": raw_text, "cleaned": cleaned_text, "summary": page_summary,
                "img_key": image_key, "json": json.dumps(vlm_result) if vlm_result else None,
            },
        )

        pages_data.append({
            "page_number": page_num + 1,
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "summary": page_summary,
            "vlm_result": vlm_result,
        })

    session.commit()
    doc.close()
    return pages_data


def _process_image(session, document_id: str, project_id: str, file_data: bytes, file_type: str) -> list[dict]:
    """Process a single image file."""
    from sqlalchemy import text

    image_key = f"projects/{project_id}/documents/{document_id}/pages/1.png"
    upload_bytes(image_key, file_data, content_type=f"image/{file_type}")

    vlm_result = _try_vlm_analysis(file_data, "")
    summary = vlm_result.get("summary", "Image document")

    page_id = str(uuid.uuid4())
    session.execute(
        text("""INSERT INTO document_pages
            (id, document_id, page_number, raw_text, cleaned_text, page_summary,
             image_storage_key, extracted_json, created_at, updated_at)
            VALUES (:id, :doc_id, 1, '', :summary, :summary, :img_key, :json, NOW(), NOW())"""),
        {
            "id": page_id, "doc_id": document_id,
            "summary": summary, "img_key": image_key,
            "json": json.dumps(vlm_result) if vlm_result else None,
        },
    )
    session.commit()

    return [{"page_number": 1, "raw_text": "", "cleaned_text": summary, "summary": summary, "vlm_result": vlm_result}]


def _process_docx(session, document_id: str, project_id: str, file_data: bytes) -> list[dict]:
    """Process a DOCX file by extracting text from paragraphs and tables."""
    from sqlalchemy import text as sql_text
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_data))
    full_text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text_parts.append(row_text)

    full_text = "\n".join(full_text_parts)

    # Split into virtual pages of ~3000 chars
    pages_data = []
    page_size = 3000
    for i in range(0, max(len(full_text), 1), page_size):
        page_num = (i // page_size) + 1
        page_text = full_text[i:i + page_size]
        if not page_text.strip():
            continue

        page_id = str(uuid.uuid4())
        session.execute(
            sql_text("""INSERT INTO document_pages
                (id, document_id, page_number, raw_text, cleaned_text, page_summary,
                 image_storage_key, extracted_json, created_at, updated_at)
                VALUES (:id, :doc_id, :pn, :raw, :cleaned, :summary, NULL, NULL, NOW(), NOW())"""),
            {
                "id": page_id, "doc_id": document_id, "pn": page_num,
                "raw": page_text, "cleaned": page_text,
                "summary": page_text[:500],
            },
        )
        pages_data.append({
            "page_number": page_num,
            "raw_text": page_text,
            "cleaned_text": page_text,
            "summary": page_text[:500],
            "vlm_result": {},
        })

    if not pages_data:
        pages_data = [{"page_number": 1, "raw_text": "", "cleaned_text": "Empty document", "summary": "Empty document", "vlm_result": {}}]

    session.commit()
    return pages_data


def _process_xlsx(session, document_id: str, project_id: str, file_data: bytes) -> list[dict]:
    """Process an XLSX file by extracting text from all sheets."""
    from sqlalchemy import text as sql_text
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_data), read_only=True, data_only=True)
    pages_data = []
    page_num = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        header = None

        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            cell_values = [str(c) if c is not None else "" for c in row]
            if not any(v.strip() for v in cell_values):
                continue
            if header is None:
                header = cell_values
            row_str = " | ".join(cell_values)
            rows_text.append(row_str)

        if not rows_text:
            continue

        # Split large sheets into chunks of ~50 rows
        chunk_size = 50
        for i in range(0, len(rows_text), chunk_size):
            page_num += 1
            chunk_rows = rows_text[i:i + chunk_size]
            page_text = f"Sheet: {sheet_name}\n" + "\n".join(chunk_rows)

            page_id = str(uuid.uuid4())
            session.execute(
                sql_text("""INSERT INTO document_pages
                    (id, document_id, page_number, raw_text, cleaned_text, page_summary,
                     image_storage_key, extracted_json, created_at, updated_at)
                    VALUES (:id, :doc_id, :pn, :raw, :cleaned, :summary, NULL, NULL, NOW(), NOW())"""),
                {
                    "id": page_id, "doc_id": document_id, "pn": page_num,
                    "raw": page_text, "cleaned": page_text,
                    "summary": f"Sheet '{sheet_name}' rows {i+1}-{i+len(chunk_rows)}",
                },
            )
            pages_data.append({
                "page_number": page_num,
                "raw_text": page_text,
                "cleaned_text": page_text,
                "summary": f"Sheet '{sheet_name}' rows {i+1}-{i+len(chunk_rows)}",
                "vlm_result": {},
            })

    wb.close()

    if not pages_data:
        pages_data = [{"page_number": 1, "raw_text": "", "cleaned_text": "Empty spreadsheet", "summary": "Empty spreadsheet", "vlm_result": {}}]

    session.commit()
    return pages_data


def _process_csv(session, document_id: str, project_id: str, file_data: bytes) -> list[dict]:
    """Process a CSV file by extracting rows as text."""
    from sqlalchemy import text as sql_text

    text_content = file_data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text_content))
    rows_text = []
    for row in reader:
        row_str = " | ".join(row)
        if row_str.strip():
            rows_text.append(row_str)

    pages_data = []
    chunk_size = 50
    for i in range(0, max(len(rows_text), 1), chunk_size):
        page_num = (i // chunk_size) + 1
        chunk_rows = rows_text[i:i + chunk_size]
        page_text = "\n".join(chunk_rows)
        if not page_text.strip():
            continue

        page_id = str(uuid.uuid4())
        session.execute(
            sql_text("""INSERT INTO document_pages
                (id, document_id, page_number, raw_text, cleaned_text, page_summary,
                 image_storage_key, extracted_json, created_at, updated_at)
                VALUES (:id, :doc_id, :pn, :raw, :cleaned, :summary, NULL, NULL, NOW(), NOW())"""),
            {
                "id": page_id, "doc_id": document_id, "pn": page_num,
                "raw": page_text, "cleaned": page_text,
                "summary": f"CSV rows {i+1}-{i+len(chunk_rows)}",
            },
        )
        pages_data.append({
            "page_number": page_num,
            "raw_text": page_text,
            "cleaned_text": page_text,
            "summary": f"CSV rows {i+1}-{i+len(chunk_rows)}",
            "vlm_result": {},
        })

    if not pages_data:
        pages_data = [{"page_number": 1, "raw_text": "", "cleaned_text": "Empty CSV", "summary": "Empty CSV", "vlm_result": {}}]

    session.commit()
    return pages_data


def _try_vlm_analysis(image_bytes: bytes, raw_text: str) -> dict:
    """Try to call VLM for page understanding. Returns empty dict on failure.

    Supports both Anthropic (Claude) and OpenAI-compatible endpoints based on
    the LLM_PROVIDER env var.
    """
    import base64
    import json

    provider = os.environ.get("LLM_PROVIDER", "anthropic")
    b64 = base64.b64encode(image_bytes).decode()
    prompt = "Analyze this construction document page. Extract: page_type, title, summary, metadata (sheet_number, revision, date, trade), entities. Return as JSON."
    if raw_text:
        prompt += f"\n\nOCR text: {raw_text[:1500]}"

    try:
        if provider == "anthropic":
            api_key = os.environ.get("ANTHROPIC_API_KEY", "")
            model = os.environ.get("ANTHROPIC_VISION_MODEL", "claude-sonnet-4-20250514")
            if not api_key:
                return {}

            payload = {
                "model": model,
                "max_tokens": 1500,
                "temperature": 0.1,
                "messages": [
                    {"role": "user", "content": [
                        {"type": "image", "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": b64,
                        }},
                        {"type": "text", "text": prompt},
                    ]},
                ],
            }

            with httpx.Client(timeout=120) as client:
                resp = client.post(
                    "https://api.anthropic.com/v1/messages",
                    json=payload,
                    headers={
                        "x-api-key": api_key,
                        "anthropic-version": "2023-06-01",
                        "Content-Type": "application/json",
                    },
                )
                resp.raise_for_status()
                data = resp.json()
                content = "".join(
                    block["text"] for block in data["content"] if block["type"] == "text"
                )
        else:
            vlm_base = os.environ.get("VLM_BASE_URL", "")
            vlm_key = os.environ.get("VLM_API_KEY", "not-needed")
            vlm_model = os.environ.get("VLM_MODEL", "gpt-4o")
            if not vlm_base:
                return {}

            payload = {
                "model": vlm_model,
                "messages": [
                    {"role": "user", "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                        {"type": "text", "text": prompt},
                    ]},
                ],
                "temperature": 0.1,
                "max_tokens": 1500,
            }

            with httpx.Client(timeout=120) as client:
                resp = client.post(
                    f"{vlm_base}/chat/completions",
                    json=payload,
                    headers={"Authorization": f"Bearer {vlm_key}", "Content-Type": "application/json"},
                )
                resp.raise_for_status()
                content = resp.json()["choices"][0]["message"]["content"]

        try:
            # Strip markdown code fences if present
            cleaned = content.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned[3:]
                if cleaned.endswith("```"):
                    cleaned = cleaned[:-3]
                cleaned = cleaned.strip()
            return json.loads(cleaned)
        except json.JSONDecodeError:
            return {"summary": content[:500], "page_type": "text"}

    except Exception as e:
        logger.debug(f"VLM analysis skipped: {e}")
        return {}


def _chunk_pages(pages_data: list[dict], doc_type: str) -> list[dict]:
    """Type-aware chunking based on document type."""
    chunks = []

    if doc_type == "drawing":
        # For drawings: one chunk per page
        for page in pages_data:
            text = page["cleaned_text"] or page["summary"] or ""
            if text.strip():
                chunks.append({
                    "page_number": page["page_number"],
                    "text": text,
                    "chunk_type": "drawing_page",
                })

    elif doc_type == "specification":
        # Section-based chunking
        for page in pages_data:
            text = page["cleaned_text"] or ""
            if not text.strip():
                continue
            sections = _split_by_sections(text)
            for i, section in enumerate(sections):
                if section.strip():
                    chunks.append({
                        "page_number": page["page_number"],
                        "text": section,
                        "chunk_type": "spec_section",
                    })

    elif doc_type in ("rfi", "submittal"):
        # Record-level chunking
        full_text = "\n\n".join(p["cleaned_text"] or p["summary"] or "" for p in pages_data)
        if full_text.strip():
            chunks.append({
                "page_number": 1,
                "text": full_text[:3000],
                "chunk_type": f"{doc_type}_record",
            })

    elif doc_type == "daily_report":
        # Date-section based
        for page in pages_data:
            text = page["cleaned_text"] or ""
            if text.strip():
                sections = _split_by_paragraphs(text, max_size=800)
                for section in sections:
                    chunks.append({
                        "page_number": page["page_number"],
                        "text": section,
                        "chunk_type": "daily_report_section",
                    })

    else:
        # General paragraph-based chunking
        for page in pages_data:
            text = page["cleaned_text"] or page["summary"] or ""
            if not text.strip():
                continue
            page_chunks = _split_by_paragraphs(text, max_size=800, overlap=100)
            for chunk_text in page_chunks:
                chunks.append({
                    "page_number": page["page_number"],
                    "text": chunk_text,
                    "chunk_type": "general",
                })

    return chunks


def _split_by_sections(text: str) -> list[str]:
    """Split text by section headers (numbered sections like 1.1, 2.3, etc.)."""
    import re
    sections = re.split(r'\n(?=\d+\.\d+[\s.])', text)
    return [s for s in sections if s.strip()]


def _split_by_paragraphs(text: str, max_size: int = 800, overlap: int = 100) -> list[str]:
    """Split text into overlapping chunks."""
    if len(text) <= max_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + max_size
        if end < len(text):
            # Try to break at paragraph
            break_point = text.rfind("\n\n", start, end)
            if break_point == -1 or break_point <= start:
                break_point = text.rfind(". ", start, end)
            if break_point > start:
                end = break_point + 1
        chunks.append(text[start:end].strip())
        start = end - overlap

    return [c for c in chunks if c.strip()]


def _embed_and_store(
    session, document_id: str, project_id: str, chunks: list[dict],
    visibility_scope: str, trade_scope: str, doc_type: str,
):
    """Generate embeddings and store in Qdrant + database."""
    from sqlalchemy import text

    embed_base = os.environ.get("EMBEDDING_BASE_URL", "")
    embed_key = os.environ.get("EMBEDDING_API_KEY", "not-needed")
    embed_model = os.environ.get("EMBEDDING_MODEL", "BAAI/bge-m3")

    texts = [c["text"] for c in chunks]

    # Get embeddings
    try:
        embeddings = _get_embeddings(texts, embed_base, embed_key, embed_model)
    except Exception as e:
        logger.error(f"Embedding failed: {e}")
        # Store chunks without embeddings
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            session.execute(
                text("""INSERT INTO document_chunks
                    (id, document_id, page_number, chunk_index, chunk_text, metadata_json,
                     visibility_scope, trade_scope, vector_id, created_at, updated_at)
                    VALUES (:id, :doc_id, :pn, :ci, :text, :meta, :vis, :trade, NULL, NOW(), NOW())"""),
                {
                    "id": chunk_id, "doc_id": document_id,
                    "pn": chunk["page_number"], "ci": i,
                    "text": chunk["text"],
                    "meta": f'{{"chunk_type": "{chunk["chunk_type"]}"}}',
                    "vis": visibility_scope, "trade": trade_scope,
                },
            )
        session.commit()
        return

    # Prepare Qdrant points
    collection_name = ensure_collection(project_id)
    points = []
    for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        chunk_id = str(uuid.uuid4())
        vector_id = chunk_id

        payload = {
            "document_id": document_id,
            "project_id": project_id,
            "page_number": chunk["page_number"],
            "chunk_index": i,
            "text": chunk["text"],
            "visibility_scope": visibility_scope,
            "trade_scope": trade_scope,
            "doc_type": doc_type,
        }

        points.append(
            qdrant_models.PointStruct(
                id=vector_id,
                vector={"dense": embedding},
                payload=payload,
            )
        )

        # Store chunk record in DB
        session.execute(
            text("""INSERT INTO document_chunks
                (id, document_id, page_number, chunk_index, chunk_text, metadata_json,
                 visibility_scope, trade_scope, vector_id, created_at, updated_at)
                VALUES (:id, :doc_id, :pn, :ci, :text, :meta, :vis, :trade, :vid, NOW(), NOW())"""),
            {
                "id": chunk_id, "doc_id": document_id,
                "pn": chunk["page_number"], "ci": i,
                "text": chunk["text"],
                "meta": f'{{"chunk_type": "{chunk["chunk_type"]}"}}',
                "vis": visibility_scope, "trade": trade_scope,
                "vid": vector_id,
            },
        )

    # Upsert to Qdrant
    upsert_points(collection_name, points)
    session.commit()
    logger.info(f"Stored {len(points)} vectors for document {document_id}")


def _get_embeddings(texts: list[str], base_url: str, api_key: str, model: str) -> list[list[float]]:
    """Get embeddings — uses local sentence-transformers or API based on EMBEDDING_PROVIDER."""
    provider = os.environ.get("EMBEDDING_PROVIDER", "local")

    if provider == "local":
        return _get_local_embeddings(texts, model)
    else:
        return _get_api_embeddings(texts, base_url, api_key, model)


def _get_local_embeddings(texts: list[str], model_name: str) -> list[list[float]]:
    """Get embeddings using fastembed (ONNX, lightweight)."""
    from fastembed import TextEmbedding

    model = TextEmbedding(model_name=model_name)
    embeddings = list(model.embed(texts))
    return [emb.tolist() for emb in embeddings]


def _get_api_embeddings(texts: list[str], base_url: str, api_key: str, model: str) -> list[list[float]]:
    """Get embeddings from an OpenAI-compatible API."""
    if not base_url:
        raise ValueError("EMBEDDING_BASE_URL not configured")

    all_embeddings = []
    batch_size = 32

    with httpx.Client(timeout=60) as client:
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            resp = client.post(
                f"{base_url}/embeddings",
                json={"model": model, "input": batch},
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            )
            resp.raise_for_status()
            data = resp.json()
            all_embeddings.extend([item["embedding"] for item in data["data"]])

    return all_embeddings
