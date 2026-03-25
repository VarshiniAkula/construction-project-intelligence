"""Document ingestion pipeline — runs as BackgroundTask (no Celery needed)."""
import io
import csv
import json
import uuid
import logging
import os

import httpx

from app.config import settings
from app.services.storage import upload_file, download_file

logger = logging.getLogger(__name__)

# Sync Supabase client for background tasks
_sb_client = None


def _get_sb():
    global _sb_client
    if _sb_client is None:
        from supabase import create_client
        _sb_client = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_KEY)
    return _sb_client


def _execute_sql(sql: str, params: dict | None = None):
    """Execute raw SQL via Supabase's REST SQL endpoint (for pgvector operations)."""
    # Use the postgrest RPC or direct SQL endpoint
    # Supabase doesn't have a direct SQL endpoint via REST, so we'll use a workaround
    # We'll create chunks without embeddings via PostgREST, then update via RPC
    pass


def ingest_document(document_id: str):
    """Full ingestion pipeline for a document. Runs synchronously in a thread."""
    sb = _get_sb()

    try:
        result = sb.table("documents").select("*").eq("id", document_id).maybe_single().execute()
        doc_row = result.data if result else None
        if not doc_row:
            logger.error(f"Document {document_id} not found")
            return

        project_id = doc_row["project_id"]
        storage_key = doc_row["storage_key"]
        file_type = storage_key.rsplit(".", 1)[-1] if "." in storage_key else "pdf"
        doc_type = doc_row["doc_type"]
        visibility_scope = doc_row["visibility_scope"]
        trade_scope = doc_row.get("trade_scope") or ""

        logger.info(f"Starting ingestion for document {document_id} ({file_type})")

        sb.table("documents").update({"status": "rendering_pages"}).eq("id", document_id).execute()

        # Download file
        file_data = download_file(storage_key)

        # Process based on type
        if file_type == "pdf":
            pages_data = _process_pdf(sb, document_id, project_id, file_data)
        elif file_type in ("png", "jpg", "jpeg"):
            pages_data = _process_image(sb, document_id, project_id, file_data, file_type)
        elif file_type == "docx":
            pages_data = _process_docx(sb, document_id, file_data)
        elif file_type in ("xlsx", "xls"):
            pages_data = _process_xlsx(sb, document_id, file_data)
        elif file_type == "csv":
            pages_data = _process_csv(sb, document_id, file_data)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Update page count
        sb.table("documents").update({"page_count": len(pages_data), "status": "chunking"}).eq("id", document_id).execute()

        # Chunk
        chunks = _chunk_pages(pages_data, doc_type)
        logger.info(f"Created {len(chunks)} chunks for document {document_id}")

        # Embed and store
        sb.table("documents").update({"status": "embedding"}).eq("id", document_id).execute()

        if chunks:
            _embed_and_store(sb, document_id, project_id, chunks, visibility_scope, trade_scope, doc_type)

        sb.table("documents").update({"status": "ready"}).eq("id", document_id).execute()
        logger.info(f"Document {document_id} ingestion complete")

    except Exception as e:
        logger.error(f"Ingestion failed for {document_id}: {e}", exc_info=True)
        try:
            sb.table("documents").update({"status": "error", "processing_error": str(e)[:1000]}).eq("id", document_id).execute()
        except Exception:
            pass


def _process_pdf(sb, document_id: str, project_id: str, file_data: bytes) -> list[dict]:
    import fitz

    doc = fitz.open(stream=file_data, filetype="pdf")
    pages_data = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        raw_text = page.get_text("text")

        # Render page image
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")

        image_key = f"projects/{project_id}/documents/{document_id}/pages/{page_num + 1}.png"
        upload_file(image_key, img_bytes, content_type="image/png")

        page_summary = raw_text[:500] if raw_text else ""
        cleaned_text = raw_text.strip()

        page_id = str(uuid.uuid4())
        sb.table("document_pages").insert({
            "id": page_id,
            "document_id": document_id,
            "page_number": page_num + 1,
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "page_summary": page_summary,
            "image_storage_key": image_key,
        }).execute()

        pages_data.append({
            "page_number": page_num + 1,
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "summary": page_summary,
        })

    doc.close()
    return pages_data


def _process_image(sb, document_id: str, project_id: str, file_data: bytes, file_type: str) -> list[dict]:
    image_key = f"projects/{project_id}/documents/{document_id}/pages/1.png"
    upload_file(image_key, file_data, content_type=f"image/{file_type}")

    page_id = str(uuid.uuid4())
    sb.table("document_pages").insert({
        "id": page_id,
        "document_id": document_id,
        "page_number": 1,
        "raw_text": "",
        "cleaned_text": "Image document",
        "page_summary": "Image document",
        "image_storage_key": image_key,
    }).execute()
    return [{"page_number": 1, "raw_text": "", "cleaned_text": "Image document", "summary": "Image document"}]


def _process_docx(sb, document_id: str, file_data: bytes) -> list[dict]:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_data))
    full_text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text_parts.append(row_text)

    full_text = "\n".join(full_text_parts)
    pages_data = []
    page_size = 3000
    for i in range(0, max(len(full_text), 1), page_size):
        page_num = (i // page_size) + 1
        page_text = full_text[i:i + page_size]
        if not page_text.strip():
            continue

        page_id = str(uuid.uuid4())
        sb.table("document_pages").insert({
            "id": page_id,
            "document_id": document_id,
            "page_number": page_num,
            "raw_text": page_text,
            "cleaned_text": page_text,
            "page_summary": page_text[:500],
        }).execute()
        pages_data.append({
            "page_number": page_num, "raw_text": page_text,
            "cleaned_text": page_text, "summary": page_text[:500],
        })

    if not pages_data:
        pages_data = [{"page_number": 1, "raw_text": "", "cleaned_text": "Empty document", "summary": "Empty document"}]

    return pages_data


def _process_xlsx(sb, document_id: str, file_data: bytes) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_data), read_only=True, data_only=True)
    pages_data = []
    page_num = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []

        for row in ws.iter_rows(values_only=True):
            cell_values = [str(c) if c is not None else "" for c in row]
            if not any(v.strip() for v in cell_values):
                continue
            rows_text.append(" | ".join(cell_values))

        if not rows_text:
            continue

        chunk_size = 50
        for i in range(0, len(rows_text), chunk_size):
            page_num += 1
            chunk_rows = rows_text[i:i + chunk_size]
            page_text = f"Sheet: {sheet_name}\n" + "\n".join(chunk_rows)

            page_id = str(uuid.uuid4())
            sb.table("document_pages").insert({
                "id": page_id,
                "document_id": document_id,
                "page_number": page_num,
                "raw_text": page_text,
                "cleaned_text": page_text,
                "page_summary": f"Sheet '{sheet_name}' rows {i+1}-{i+len(chunk_rows)}",
            }).execute()
            pages_data.append({
                "page_number": page_num, "raw_text": page_text,
                "cleaned_text": page_text,
                "summary": f"Sheet '{sheet_name}' rows {i+1}-{i+len(chunk_rows)}",
            })

    wb.close()
    if not pages_data:
        pages_data = [{"page_number": 1, "raw_text": "", "cleaned_text": "Empty spreadsheet", "summary": "Empty spreadsheet"}]

    return pages_data


def _process_csv(sb, document_id: str, file_data: bytes) -> list[dict]:
    text_content = file_data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text_content))
    rows_text = [" | ".join(row) for row in reader if any(c.strip() for c in row)]

    pages_data = []
    chunk_size = 50
    for i in range(0, max(len(rows_text), 1), chunk_size):
        page_num = (i // chunk_size) + 1
        chunk_rows = rows_text[i:i + chunk_size]
        page_text = "\n".join(chunk_rows)
        if not page_text.strip():
            continue

        page_id = str(uuid.uuid4())
        sb.table("document_pages").insert({
            "id": page_id,
            "document_id": document_id,
            "page_number": page_num,
            "raw_text": page_text,
            "cleaned_text": page_text,
            "page_summary": f"CSV rows {i+1}-{i+len(chunk_rows)}",
        }).execute()
        pages_data.append({
            "page_number": page_num, "raw_text": page_text,
            "cleaned_text": page_text, "summary": f"CSV rows {i+1}-{i+len(chunk_rows)}",
        })

    if not pages_data:
        pages_data = [{"page_number": 1, "raw_text": "", "cleaned_text": "Empty CSV", "summary": "Empty CSV"}]

    return pages_data


def _chunk_pages(pages_data: list[dict], doc_type: str) -> list[dict]:
    chunks = []

    if doc_type == "drawing":
        for page in pages_data:
            t = page["cleaned_text"] or page["summary"] or ""
            if t.strip():
                chunks.append({"page_number": page["page_number"], "text": t, "chunk_type": "drawing_page"})

    elif doc_type == "specification":
        for page in pages_data:
            t = page["cleaned_text"] or ""
            if not t.strip():
                continue
            import re
            sections = re.split(r'\n(?=\d+\.\d+[\s.])', t)
            for section in sections:
                if section.strip():
                    chunks.append({"page_number": page["page_number"], "text": section, "chunk_type": "spec_section"})

    elif doc_type in ("rfi", "submittal"):
        full_text = "\n\n".join(p["cleaned_text"] or p["summary"] or "" for p in pages_data)
        if full_text.strip():
            chunks.append({"page_number": 1, "text": full_text[:3000], "chunk_type": f"{doc_type}_record"})

    elif doc_type == "daily_report":
        for page in pages_data:
            t = page["cleaned_text"] or ""
            if t.strip():
                for section in _split_by_paragraphs(t, max_size=800):
                    chunks.append({"page_number": page["page_number"], "text": section, "chunk_type": "daily_report_section"})

    else:
        for page in pages_data:
            t = page["cleaned_text"] or page["summary"] or ""
            if not t.strip():
                continue
            for chunk_text in _split_by_paragraphs(t, max_size=800, overlap=100):
                chunks.append({"page_number": page["page_number"], "text": chunk_text, "chunk_type": "general"})

    return chunks


def _split_by_paragraphs(t: str, max_size: int = 800, overlap: int = 100) -> list[str]:
    if len(t) <= max_size:
        return [t]
    chunks = []
    start = 0
    while start < len(t):
        end = start + max_size
        if end < len(t):
            bp = t.rfind("\n\n", start, end)
            if bp == -1 or bp <= start:
                bp = t.rfind(". ", start, end)
            if bp > start:
                end = bp + 1
        chunks.append(t[start:end].strip())
        start = end - overlap
    return [c for c in chunks if c.strip()]


def _embed_and_store(sb, document_id, project_id, chunks, visibility_scope, trade_scope, doc_type):
    """Generate embeddings and store chunks. Embeddings are stored via direct HTTP call."""
    texts = [c["text"] for c in chunks]

    # Skip embeddings on Vercel serverless (no persistent model cache)
    embeddings = None
    if not os.environ.get("VERCEL"):
        try:
            embeddings = _get_embeddings(texts)
        except Exception as e:
            logger.error(f"Embedding failed: {e}")
            embeddings = None
    else:
        logger.info("Vercel serverless: skipping embeddings, keyword search available")

    # Store chunks via PostgREST
    chunk_ids = []
    for i, chunk in enumerate(chunks):
        chunk_id = str(uuid.uuid4())
        chunk_ids.append(chunk_id)
        sb.table("document_chunks").insert({
            "id": chunk_id,
            "document_id": document_id,
            "page_number": chunk["page_number"],
            "chunk_index": i,
            "chunk_text": chunk["text"],
            "doc_type": doc_type,
            "metadata_json": {"chunk_type": chunk["chunk_type"]},
            "visibility_scope": visibility_scope,
            "trade_scope": trade_scope,
            "vector_id": chunk_id,
        }).execute()

    # Update embeddings via Supabase SQL HTTP API if we have them
    if embeddings:
        _update_embeddings_via_http(chunk_ids, embeddings)

    logger.info(f"Stored {len(chunks)} chunks for document {document_id}")


def _update_embeddings_via_http(chunk_ids: list[str], embeddings: list[list[float]]):
    """Update chunk embeddings using Supabase's pg_net or direct SQL call."""
    # Use the postgrest RPC endpoint to call a database function for vector updates
    # Since PostgREST can't handle vector type casting, we use the SQL endpoint
    url = f"{settings.SUPABASE_URL}/rest/v1/rpc/update_chunk_embedding"
    headers = {
        "apikey": settings.SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {settings.SUPABASE_SERVICE_KEY}",
        "Content-Type": "application/json",
    }

    with httpx.Client(timeout=30) as client:
        for chunk_id, embedding in zip(chunk_ids, embeddings):
            try:
                # Call RPC function to update embedding
                resp = client.post(url, json={
                    "p_chunk_id": chunk_id,
                    "p_embedding": embedding,
                }, headers=headers)
                if resp.status_code >= 400:
                    logger.warning(f"Failed to update embedding for {chunk_id}: {resp.text}")
            except Exception as e:
                logger.warning(f"Failed to update embedding for {chunk_id}: {e}")


def ingest_document_lightweight(document_id: str):
    """Lightweight ingestion designed to complete within 10s on Vercel serverless.

    Skips page image rendering and embeddings entirely.
    Does: text extraction -> page creation -> chunking -> chunk storage.
    """
    sb = _get_sb()

    try:
        result = sb.table("documents").select("*").eq("id", document_id).maybe_single().execute()
        doc_row = result.data if result else None
        if not doc_row:
            logger.error(f"Document {document_id} not found")
            return

        project_id = doc_row["project_id"]
        storage_key = doc_row["storage_key"]
        file_type = storage_key.rsplit(".", 1)[-1] if "." in storage_key else "pdf"
        doc_type = doc_row["doc_type"]
        visibility_scope = doc_row["visibility_scope"]
        trade_scope = doc_row.get("trade_scope") or ""

        logger.info(f"Starting lightweight ingestion for document {document_id} ({file_type})")
        sb.table("documents").update({"status": "processing"}).eq("id", document_id).execute()

        file_data = download_file(storage_key)

        # Process based on type — PDF uses text-only path (no pixmap rendering)
        if file_type == "pdf":
            pages_data = _process_pdf_text_only(sb, document_id, project_id, file_data)
        elif file_type in ("png", "jpg", "jpeg"):
            pages_data = _process_image(sb, document_id, project_id, file_data, file_type)
        elif file_type == "docx":
            pages_data = _process_docx(sb, document_id, file_data)
        elif file_type in ("xlsx", "xls"):
            pages_data = _process_xlsx(sb, document_id, file_data)
        elif file_type == "csv":
            pages_data = _process_csv(sb, document_id, file_data)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        sb.table("documents").update({"page_count": len(pages_data), "status": "chunking"}).eq("id", document_id).execute()

        chunks = _chunk_pages(pages_data, doc_type)
        logger.info(f"Created {len(chunks)} chunks for document {document_id}")

        # Store chunks without embeddings (keyword search only)
        if chunks:
            for i, chunk in enumerate(chunks):
                chunk_id = str(uuid.uuid4())
                sb.table("document_chunks").insert({
                    "id": chunk_id,
                    "document_id": document_id,
                    "page_number": chunk["page_number"],
                    "chunk_index": i,
                    "chunk_text": chunk["text"],
                    "doc_type": doc_type,
                    "metadata_json": {"chunk_type": chunk["chunk_type"]},
                    "visibility_scope": visibility_scope,
                    "trade_scope": trade_scope,
                    "vector_id": chunk_id,
                }).execute()

        sb.table("documents").update({"status": "ready"}).eq("id", document_id).execute()
        logger.info(f"Document {document_id} lightweight ingestion complete")

    except Exception as e:
        logger.error(f"Lightweight ingestion failed for {document_id}: {e}", exc_info=True)
        try:
            sb.table("documents").update({"status": "error", "processing_error": str(e)[:1000]}).eq("id", document_id).execute()
        except Exception:
            pass


def _process_pdf_text_only(sb, document_id: str, project_id: str, file_data: bytes) -> list[dict]:
    """Extract text from PDF without rendering page images (fast path for serverless)."""
    import fitz

    doc = fitz.open(stream=file_data, filetype="pdf")
    pages_data = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        raw_text = page.get_text("text")

        page_summary = raw_text[:500] if raw_text else ""
        cleaned_text = raw_text.strip()

        page_id = str(uuid.uuid4())
        sb.table("document_pages").insert({
            "id": page_id,
            "document_id": document_id,
            "page_number": page_num + 1,
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "page_summary": page_summary,
        }).execute()

        pages_data.append({
            "page_number": page_num + 1,
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "summary": page_summary,
        })

    doc.close()
    return pages_data


def _get_embeddings(texts: list[str]) -> list[list[float]]:
    provider = os.environ.get("EMBEDDING_PROVIDER", settings.EMBEDDING_PROVIDER)

    if provider == "local":
        from fastembed import TextEmbedding
        model_name = os.environ.get("EMBEDDING_MODEL", settings.EMBEDDING_MODEL)
        model = TextEmbedding(model_name=model_name)
        return [emb.tolist() for emb in model.embed(texts)]
    else:
        base_url = settings.EMBEDDING_BASE_URL
        api_key = settings.EMBEDDING_API_KEY
        model = settings.EMBEDDING_MODEL
        all_embeddings = []
        batch_size = 32
        with httpx.Client(timeout=60) as client:
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                resp = client.post(
                    f"{base_url}/embeddings",
                    json={"model": model, "input": batch},
                    headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                )
                resp.raise_for_status()
                data = resp.json()
                all_embeddings.extend([item["embedding"] for item in data["data"]])
        return all_embeddings
